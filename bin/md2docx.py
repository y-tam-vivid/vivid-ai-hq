#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Markdown を日本語の契約書・提出文書らしい .docx へ変換する。

なぜ在るか ── この機には pandoc も LibreOffice も入っていない（2026-08-26 実測）。
契約書は最終的に Word で相手へ渡すので、md のまま置いておくと必ず人の手が要る。

使い方:
    python3 bin/md2docx.py <入力.md> [出力.docx]

扱える記法（契約書に出るものだけ。汎用コンバータではない）:
    # 見出し1        → 表題（中央寄せ・太字）
    ## 見出し2       → 条見出し（太字）
    ### 見出し3      → 小見出し（太字・やや小）
    | a | b |        → Word の表（1行目をヘッダ扱い）
    - 項目           → 箇条書き
    1. 項目          → 番号付き
    **強調**         → 太字（行の途中でも効く）
    ---              → 区切り（空段落）

★変換した .docx は必ず Word で開いて目視すること。
  条番号の自動採番と手書き番号が二重になっていないかは、機械では見ていない。
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

JP_FONT = "游明朝"
JP_FONT_FALLBACK = "MS Mincho"


def set_font(run, size=None, bold=None):
    run.font.name = JP_FONT
    run.font.size = Pt(size) if size else Pt(10.5)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), JP_FONT)
    rfonts.set(qn("w:ascii"), JP_FONT_FALLBACK)
    rfonts.set(qn("w:hAnsi"), JP_FONT_FALLBACK)


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_rich_text(paragraph, text, size=None, base_bold=False):
    """**強調** を太字の run に分けて流し込む。"""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            set_font(paragraph.add_run(text[pos:m.start()]), size, base_bold)
        set_font(paragraph.add_run(m.group(1)), size, True)
        pos = m.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]), size, base_bold)
    if not paragraph.runs:
        set_font(paragraph.add_run(""), size, base_bold)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator_row(line):
    return bool(re.fullmatch(r"\|[\s:\-|]+\|", line.strip()))


def convert(src: Path, dst: Path):
    lines = src.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = JP_FONT
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), JP_FONT)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 表
        if stripped.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                if not is_separator_row(lines[i]):
                    rows.append(split_row(lines[i]))
                i += 1
            if rows:
                width = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=width)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, row in enumerate(rows):
                    for c_idx in range(width):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        text = row[c_idx] if c_idx < len(row) else ""
                        add_rich_text(
                            cell.paragraphs[0], text, size=9.5,
                            base_bold=(r_idx == 0),
                        )
                doc.add_paragraph()
            continue

        # 区切り
        if re.fullmatch(r"-{3,}", stripped):
            doc.add_paragraph()
            i += 1
            continue

        # 見出し
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            add_rich_text(p, stripped[4:], size=11, base_bold=True)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            add_rich_text(p, stripped[3:], size=12, base_bold=True)
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(16)
            add_rich_text(p, stripped[2:], size=15, base_bold=True)
            i += 1
            continue

        # 箇条書き
        m_ul = re.match(r"^-\s+(.*)$", stripped)
        if m_ul:
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, m_ul.group(1))
            i += 1
            continue

        m_ol = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if m_ol:
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, m_ol.group(1))
            i += 1
            continue

        # 引用は本文として扱う（契約書では条文の引用に使われるため）
        if stripped.startswith("> "):
            stripped = stripped[2:]

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_rich_text(p, stripped)
        i += 1

    doc.save(str(dst))
    return dst


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    src = Path(sys.argv[1]).expanduser()
    if not src.exists():
        print("入力が見つかりません: %s" % src)
        sys.exit(1)
    dst = Path(sys.argv[2]).expanduser() if len(sys.argv) > 2 else src.with_suffix(".docx")
    convert(src, dst)
    print("書き出しました: %s (%d バイト)" % (dst, dst.stat().st_size))


if __name__ == "__main__":
    main()
